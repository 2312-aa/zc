"""
file_parser.py — 文档解析模块
支持格式：pdf/txt/md/docx/excel
自动清洗乱码、多余换行空格
"""
import os
from pathlib import Path
from typing import List

from utils.logger_util import get_logger
from utils.exception_util import FileParseError

logger = get_logger("FileParser")


def _clean_text(text: str) -> str:
    """
    文本清洗：去除多余空格、换行、乱码字符
    Args:
        text: 原始文本
    Returns:
        清洗后的文本
    """
    if not text:
        return ""
    # 替换连续多个换行为单个换行
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 替换连续多个空格为单个空格（保留换行）
    text = re.sub(r'[^\S\n]{2,}', ' ', text)
    # 去除行首行尾空白
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(line for line in lines if line)
    # 去除常见乱码控制字符（保留换行和制表符）
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text.strip()


def parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return _clean_text('\n'.join(text_parts))
    except Exception as e:
        raise FileParseError(
            f"PDF文件解析失败: {os.path.basename(file_path)}",
            detail=str(e)
        )


def parse_txt(file_path: str) -> str:
    """解析TXT文件"""
    try:
        # 尝试多种编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                return _clean_text(text)
            except UnicodeDecodeError:
                continue
        raise FileParseError("无法识别文件编码", detail=file_path)
    except FileParseError:
        raise
    except Exception as e:
        raise FileParseError(
            f"TXT文件解析失败: {os.path.basename(file_path)}",
            detail=str(e)
        )


def parse_md(file_path: str) -> str:
    """解析Markdown文件（按纯文本处理）"""
    return parse_txt(file_path)


def parse_docx(file_path: str) -> str:
    """解析DOCX文件"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        return _clean_text('\n'.join(text_parts))
    except Exception as e:
        raise FileParseError(
            f"DOCX文件解析失败: {os.path.basename(file_path)}",
            detail=str(e)
        )


def parse_excel(file_path: str) -> str:
    """解析Excel文件（xlsx/xls）"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"[工作表: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                row_text = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        wb.close()
        return _clean_text('\n'.join(text_parts))
    except Exception as e:
        raise FileParseError(
            f"Excel文件解析失败: {os.path.basename(file_path)}",
            detail=str(e)
        )


# 文件扩展名与解析函数的映射
PARSER_MAP = {
    ".pdf": parse_pdf,
    ".txt": parse_txt,
    ".md": parse_md,
    ".docx": parse_docx,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
}


def parse_file(file_path: str) -> str:
    """
    根据文件扩展名自动选择解析器
    Args:
        file_path: 文件路径
    Returns:
        解析后的纯文本内容
    Raises:
        FileParseError: 解析失败
    """
    ext = Path(file_path).suffix.lower()
    parser = PARSER_MAP.get(ext)

    if parser is None:
        raise FileParseError(
            f"不支持的文件格式: {ext}",
            detail=f"支持的格式: {list(PARSER_MAP.keys())}"
        )

    logger.info(f"开始解析文件: {os.path.basename(file_path)} (格式: {ext})")
    text = parser(file_path)
    logger.info(f"文件解析完成: {os.path.basename(file_path)}, 文本长度: {len(text)}")
    return text
